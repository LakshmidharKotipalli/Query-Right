import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document


def load_pdf(file_path: str) -> List[Document]:
    """Load PDF preserving page numbers in metadata."""
    import fitz  # PyMuPDF

    docs = []
    pdf = fitz.open(file_path)
    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text()
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={
                    "source": os.path.basename(file_path),
                    "page": page_num,
                    "file_path": file_path,
                    "file_type": "pdf",
                },
            ))
    pdf.close()
    return docs


def load_txt(file_path: str) -> List[Document]:
    """Load plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(
        page_content=text,
        metadata={
            "source": os.path.basename(file_path),
            "page": 1,
            "file_path": file_path,
            "file_type": "txt",
        },
    )]


def load_docx(file_path: str) -> List[Document]:
    """Load DOCX preserving paragraph structure."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = "\n".join(
        para.text for para in doc.paragraphs if para.text.strip()
    )
    return [Document(
        page_content=full_text,
        metadata={
            "source": os.path.basename(file_path),
            "page": 1,
            "file_path": file_path,
            "file_type": "docx",
        },
    )]


LOADERS = {
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".docx": load_docx,
}


def load_document(file_path: str) -> List[Document]:
    """Dispatch to appropriate loader based on file extension."""
    ext = Path(file_path).suffix.lower()
    loader_fn = LOADERS.get(ext)
    if loader_fn is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(LOADERS.keys())}")
    return loader_fn(file_path)
